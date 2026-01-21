from docx import Document
from docx.shared import Pt
import requests
from bs4 import BeautifulSoup
import time
import os

doc = Document()
os.makedirs("주독일대사관", exist_ok=True)
BASE_ID = 2975116
BASE_URL = "https://de.mofa.go.kr/de-ko/brd/m_7204/view.do?seq="
COUNT = 10
target_urls = []
for num in range(0, COUNT):
    # print(f"{BASE_URL}{BASE_ID - num}&page=1")
    target_urls.append(f"{BASE_URL}{BASE_ID - num}")

for url in target_urls:
    response = requests.get(url)

    # 통신성공여부
    # 200이면 통신은 성공한것
    # print(response.status_code)

    soup = BeautifulSoup(response.text, "html.parser")

    # 게시글 제목, 작성자, 작성일
    # select는 언제나 리스트로 리턴
    # 하나만 가져오고 싶을 때 select_one
    # 위에서 첫번째 발견된 bo_head 클래스를 가진 태그
    header_tag = soup.select_one(".bo_head")

    # 게시글 누락 가능성존재 (패턴에서 벗어나는 경우)
    if header_tag is None:
        print(" -> 에러페이지 발생!! 건너뜁니다.\n")
        continue

    # 타이틀 추출
    title_tag = header_tag.select_one("h2")
    title = title_tag.text.strip()
    text_title = doc.add_heading(f"제목: {title}\n\n", level=1)
    text_title.runs[0].font.name = "Malgun Gothic"

    # 작성자, 작성일 -> dl 태그 안에 dd 태그들
    dd_tags = header_tag.select("dl dd")
    # 저자 추출
    author_tag = dd_tags[0]
    author = author_tag.text.strip()
    text_author = doc.add_paragraph(f"저자: {author}\n")
    text_author.runs[0].font.name = "Malgun Gothic"
    # 작성일 추출
    date_tag = dd_tags[1]
    date = date_tag.text.strip()
    text_date = doc.add_paragraph(f"작성일: {date}\n")
    text_date.runs[0].font.name = "Malgun Gothic"

    bo_con_tags = soup.find_all("div", class_="bo_con")
    # contents_tag = soup.select_one(".bo_con")
    # contents = contents_tag.text.strip()
    for bo_con_tag in bo_con_tags:
        text_tag = bo_con_tag.find("div", class_="se-contents")
        text = text_tag.text
        doc.add_paragraph(text)

    time.sleep(1.5)

    doc.save("./주독일대사관/주독일대사관1.docx")