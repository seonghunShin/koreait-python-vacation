# 파이썬 -> docx(워드파일) 생성
# pip install python-docx
from docx import Document
# 기본폰트가 영어로 지정되어 있어서 한글은 깨짐
from docx.shared import Pt

# 빈 워드문서 생성
doc = Document()

# 제목 추가하기
# 제목크기 - level 1, 2, 3
title = doc.add_heading("파이썬스터디", level = 1)
title.runs[0].font.name = "Malgun Gothic"

# 한 줄 추가하기
pg1 = doc.add_paragraph("파이썬 공부하고 있습니다.")
pg1.runs[0].font.name = "Malgun Gothic"

doc.add_paragraph("") # 엔터
pg1 = doc.add_paragraph("열심히 고고")
pg1.runs[0].font.name = "Malgun Gothic"

# 파일 저장
# save("저장 경로")
doc.save("./파이썬테스트.docx")


##### os 모듈 #####
# 파일이나 폴더를 파이썬으로 다룰 때
import os

# 1. 파일이나 폴더가 존재하는지?
# os.path.exists(경로)
if os.path.exists("./파이썬테스트.docx"):
    print("파이썬테스트 파일 존재합니다!")

if os.path.exists("./자바테스트.docx"):
    print("자바테스트 파일 존재안합니다.")

# 2. 폴더 만들기
# makedirs()
# exist_ok -> 이미 있어도 괜찮아? 네
# 이미 같은 이름의 폴더가 존재해도 에러 x
os.makedirs("내가만든폴더", exist_ok=True)

## 도전!) 제목: 내가만든 워드, 글: "헬로 파이썬"
# 을 내가만든폴더2에 저장해보세요.

os.makedirs("내가만든폴더2", exist_ok=True)
doc2 = Document()
title2 = doc2.add_heading("내가만든 워드", level=1)
title2.runs[0].font.name = "Malgun Gothic"
pg2 = doc2.add_paragraph("헬로 파이썬!")
pg2.runs[0].font.name = "Malgun Gothic"
doc2.save("./내가만든폴더2/내가만든 워드.docx")
